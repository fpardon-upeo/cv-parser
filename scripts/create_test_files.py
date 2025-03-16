"""
Script to create test PDF and DOCX files from the sample resume text.
"""
import os
import io
from PyPDF2 import PdfWriter, PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_test_pdf(content, output_path):
    """Create a test PDF file with the given content using reportlab."""
    # Create a memory buffer for the PDF
    buffer = io.BytesIO()
    
    # Create the PDF with reportlab
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Split content into lines
    lines = content.split('\n')
    y_position = height - 50  # Start from top
    
    for line in lines:
        if line.strip():
            c.drawString(50, y_position, line)
        y_position -= 15  # Move down for next line
        
        # If we run out of space, create a new page
        if y_position < 50:
            c.showPage()
            y_position = height - 50
    
    c.save()
    
    # Move to the beginning of the buffer
    buffer.seek(0)
    
    # Create a PdfReader object from the buffer
    reader = PdfReader(buffer)
    
    # Create a PdfWriter object
    writer = PdfWriter()
    
    # Add all pages from the reader to the writer
    for page in reader.pages:
        writer.add_page(page)
    
    # Add metadata
    writer.add_metadata({
        '/Title': 'Sample Resume',
        '/Author': 'John Doe',
        '/Subject': 'Resume',
        '/Keywords': 'resume,software engineer,python'
    })
    
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Write to file
    with open(output_path, 'wb') as f:
        writer.write(f)
    
    print(f'Created PDF file at {output_path}')


def create_test_docx(content, output_path):
    """Create a test DOCX file with the given content."""
    doc = Document()
    
    # Split content into lines
    lines = content.split('\n')
    
    # Process lines
    for line in lines:
        if not line.strip():
            # Empty line
            doc.add_paragraph()
        elif line.strip().isupper() and len(line.strip()) > 3:
            # Section header (all caps)
            p = doc.add_paragraph(line)
            p.style = 'Heading 1'
        elif ':' in line and line.split(':')[0].strip().istitle():
            # Field with label (e.g., "Skills: Python")
            p = doc.add_paragraph()
            parts = line.split(':', 1)
            p.add_run(parts[0] + ':').bold = True
            if len(parts) > 1:
                p.add_run(parts[1])
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Add metadata
    doc.core_properties.title = "Sample Resume"
    doc.core_properties.author = "John Doe"
    doc.core_properties.subject = "Resume"
    doc.core_properties.keywords = "resume,software engineer,python"
    
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Save the document
    doc.save(output_path)
    
    print(f'Created DOCX file at {output_path}')


def main():
    """Main function to create test files."""
    # Read the content from the TXT file
    txt_path = 'tests/test_data/sample_resume.txt'
    pdf_path = 'tests/test_data/sample_resume.pdf'
    docx_path = 'tests/test_data/sample_resume.docx'
    
    try:
        with open(txt_path, 'r') as f:
            content = f.read()
        
        # Create the PDF file
        create_test_pdf(content, pdf_path)
        
        # Create the DOCX file
        create_test_docx(content, docx_path)
        
        print("Test files created successfully!")
    except Exception as e:
        print(f"Error creating test files: {e}")


if __name__ == "__main__":
    main() 