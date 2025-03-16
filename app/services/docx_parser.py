"""DOCX document parsing service.

This module provides the implementation for parsing DOCX documents.
"""
import io
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from datetime import datetime
from typing import Dict, Any, List, Optional

from app.services.document_parser import (
    DocumentParser,
    DocumentParsingError,
    ValidationError,
    DocumentParserFactory
)


class DOCXParsingError(DocumentParsingError):
    """Raised when DOCX parsing fails."""
    pass


class DOCXCorruptedError(DocumentParsingError):
    """Raised when a corrupted DOCX is encountered."""
    pass


class DOCXParser(DocumentParser):
    """Parser for DOCX documents."""
    
    def __init__(self):
        super().__init__()
        self._doc = None
        self._structure = None
    
    def parse(self, content: bytes) -> Dict[str, Any]:
        """Parse a DOCX document and extract its content and structure.
        
        Args:
            content: Raw DOCX file content as bytes
            
        Returns:
            Dict containing parsed content, structure, and metadata
            
        Raises:
            DOCXParsingError: If any parsing error occurs
            DOCXCorruptedError: If the DOCX is corrupted
            ValidationError: If the content fails validation
        """
        try:
            self._set_content(content)
            file_stream = io.BytesIO(content)
            self._doc = Document(file_stream)
            
            parsed_content = self._extract_text()
            structure = self._extract_structure()
            metadata = self.extract_metadata()
            
            return {
                "content": parsed_content,
                "structure": structure,
                "metadata": metadata
            }
            
        except PackageNotFoundError as e:
            raise DOCXCorruptedError(f"Failed to read DOCX: {str(e)}")
        except (ValueError, TypeError, AttributeError) as e:
            if "not a zip file" in str(e).lower():
                raise DOCXCorruptedError(f"Invalid DOCX format: {str(e)}")
            raise DOCXParsingError(f"DOCX parsing failed: {str(e)}")
        except Exception as e:
            raise DOCXParsingError(f"DOCX parsing failed: {str(e)}")
    
    def extract_metadata(self) -> Dict[str, Any]:
        """Extract document metadata.
        
        Returns:
            Dict containing document metadata
            
        Raises:
            DOCXParsingError: If metadata extraction fails
        """
        if not self._doc:
            raise DOCXParsingError("No document loaded for metadata extraction")
            
        try:
            props = self._doc.core_properties
            metadata = {}
            
            if props.title:
                metadata['title'] = props.title
            if props.author:
                metadata['author'] = props.author
            if props.subject:
                metadata['subject'] = props.subject
            if props.keywords:
                metadata['keywords'] = props.keywords
            if props.category:
                metadata['category'] = props.category
            if props.comments:
                metadata['comments'] = props.comments
            
            # Handle dates
            if props.created:
                metadata['creation_date'] = self._format_datetime(props.created)
            if props.modified:
                metadata['modification_date'] = self._format_datetime(props.modified)
            if props.last_printed:
                metadata['last_printed'] = self._format_datetime(props.last_printed)
            
            # Add document structure information
            metadata['paragraph_count'] = len(self._doc.paragraphs)
            metadata['table_count'] = len(self._doc.tables)
            metadata['section_count'] = len(self._doc.sections)
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX metadata: {str(e)}")
            return {}
    
    def validate(self) -> bool:
        """Validate DOCX content.
        
        Returns:
            bool: True if document is valid
            
        Raises:
            ValidationError: If validation fails
        """
        if not self._get_content():
            raise DOCXParsingError("No content has been set for validation")
            
        try:
            # Check if content starts with DOCX signature (PK zip header)
            if not self._content.startswith(b'PK'):
                raise ValidationError("Not a valid DOCX file (missing ZIP signature)")
                
            # Check if we can read the DOCX
            file_stream = io.BytesIO(self._content)
            doc = Document(file_stream)
            
            # Check if document has content
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                self.logger.warning("DOCX document has no paragraphs or tables")
                
            return True
            
        except PackageNotFoundError as e:
            raise ValidationError(f"Invalid DOCX structure: {str(e)}")
        except Exception as e:
            raise ValidationError(f"DOCX validation failed: {str(e)}")
    
    def _extract_text(self) -> str:
        """Extract text content from DOCX.
        
        Returns:
            str: Extracted text content
            
        Raises:
            DOCXParsingError: If text extraction fails
        """
        if not self._doc:
            raise DOCXParsingError("No document loaded for text extraction")
            
        try:
            text_content = []
            
            # Extract text from paragraphs
            for para in self._doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in self._doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise DOCXParsingError(f"Failed to extract text: {str(e)}")
    
    def _extract_structure(self) -> Dict[str, Any]:
        """Extract document structure from DOCX.
        
        Returns:
            Dict containing document structure
            
        Raises:
            DOCXParsingError: If structure extraction fails
        """
        if not self._doc:
            raise DOCXParsingError("No document loaded for structure extraction")
            
        try:
            structure = {
                'sections': [],
                'headings': [],
                'tables': [],
                'lists': []
            }
            
            # Extract section information
            for i, section in enumerate(self._doc.sections):
                section_info = {
                    'number': i + 1,
                    'page_width': section.page_width.pt,
                    'page_height': section.page_height.pt,
                    'orientation': 'portrait' if section.page_width < section.page_height else 'landscape'
                }
                structure['sections'].append(section_info)
            
            # Extract heading information (paragraphs with heading styles)
            for i, para in enumerate(self._doc.paragraphs):
                if para.style.name.startswith('Heading'):
                    heading_info = {
                        'level': int(para.style.name.replace('Heading ', '')),
                        'text': para.text,
                        'position': i
                    }
                    structure['headings'].append(heading_info)
            
            # Extract table information
            for i, table in enumerate(self._doc.tables):
                table_info = {
                    'number': i + 1,
                    'rows': len(table.rows),
                    'columns': len(table.columns)
                }
                structure['tables'].append(table_info)
            
            return structure
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX structure: {str(e)}")
            return {'sections': [], 'headings': [], 'tables': [], 'lists': []}
    
    def _format_datetime(self, dt: Optional[datetime]) -> str:
        """Format datetime object to ISO format string.
        
        Args:
            dt: datetime object or None
            
        Returns:
            str: ISO formatted date string or empty string if dt is None
        """
        return dt.isoformat() if dt else ""

# Register the DOCX parser with the factory
DocumentParserFactory.register_parser("docx", DOCXParser) 