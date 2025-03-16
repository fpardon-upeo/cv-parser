"""Unit tests for DOCX parser implementation."""
import pytest
from datetime import datetime
from docx import Document
from docx.shared import Inches
from io import BytesIO
from typing import Dict, Any

from app.services.docx_parser import (
    DOCXParser,
    DOCXParsingError,
    DOCXCorruptedError
)
from app.services.document_parser import ValidationError


def create_test_docx(content: str = "Test content", metadata: Dict[str, Any] = None) -> bytes:
    """Create a test DOCX file with given content and metadata."""
    doc = Document()
    
    # Add content
    doc.add_paragraph(content)
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Data 1"
    table.cell(1, 1).text = "Data 2"
    
    # Add metadata if provided
    if metadata:
        doc.core_properties.title = metadata.get("title", "")
        doc.core_properties.author = metadata.get("author", "")
        doc.core_properties.subject = metadata.get("subject", "")
        doc.core_properties.keywords = metadata.get("keywords", "")
    
    # Save to bytes buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_complex_docx() -> bytes:
    """Create a complex DOCX file with various elements."""
    doc = Document()
    
    # Add sections with different orientations
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    
    # Add header and footer
    header = section.header
    header.paragraphs[0].text = "Document Header"
    footer = section.footer
    footer.paragraphs[0].text = "Page Footer"
    
    # Add content with different styles
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("Regular paragraph content")
    doc.add_heading("Section 1", 1)
    doc.add_paragraph("Section content with some formatting").add_run(" (bold)").bold = True
    
    # Add table
    table = doc.add_table(rows=3, cols=3)
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f"Cell {i},{j}"
    
    # Save to bytes buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


class TestDOCXParser:
    """Test suite for DOCXParser class."""
    
    def test_parse_basic_docx(self):
        """Test parsing a basic DOCX with text content."""
        parser = DOCXParser()
        content = create_test_docx("Sample DOCX content")
        
        result = parser.parse(content)
        
        assert "content" in result
        assert "Sample DOCX content" in result["content"]
        assert "Header 1 | Header 2" in result["content"]
        assert "Data 1 | Data 2" in result["content"]
        assert "structure" in result
        assert "metadata" in result
        
    def test_parse_with_metadata(self):
        """Test parsing DOCX with metadata."""
        metadata = {
            "title": "Test Document",
            "author": "Test Author",
            "subject": "Test Subject",
            "keywords": "test, document, unit test"
        }
        
        parser = DOCXParser()
        content = create_test_docx("Content", metadata)
        
        result = parser.parse(content)
        
        assert result["metadata"]["title"] == "Test Document"
        assert result["metadata"]["author"] == "Test Author"
        assert result["metadata"]["subject"] == "Test Subject"
        assert result["metadata"]["keywords"] == "test, document, unit test"
        
    def test_parse_corrupted_docx(self):
        """Test parsing a corrupted DOCX raises appropriate error."""
        parser = DOCXParser()
        content = b"This is not a valid DOCX file"
        
        with pytest.raises(DOCXCorruptedError) as exc_info:
            parser.parse(content)
        assert "Failed to read DOCX" in str(exc_info.value)
        
    def test_extract_structure(self):
        """Test extraction of DOCX structure information."""
        parser = DOCXParser()
        content = create_complex_docx()
        
        result = parser.parse(content)
        structure = result["structure"]
        
        assert "sections" in structure
        assert len(structure["sections"]) > 0
        assert structure["sections"][0]["orientation"] == "portrait"
        assert structure["sections"][0]["page_height"] == 11
        assert structure["sections"][0]["page_width"] == 8.5
        assert structure["paragraphs"] > 0
        assert structure["tables"] > 0
        assert structure["has_headers"] is True
        assert structure["has_footers"] is True
        
    def test_validation(self):
        """Test DOCX validation."""
        parser = DOCXParser()
        content = create_test_docx()
        
        # Set content before validation
        parser._set_content(content)
        parser._doc = Document(BytesIO(content))
        
        assert parser.validate() is True
        
    def test_validation_no_content(self):
        """Test validation without content raises error."""
        parser = DOCXParser()
        
        with pytest.raises(DOCXParsingError) as exc_info:
            parser.validate()
        assert "No content to validate" in str(exc_info.value)
        
    def test_text_extraction(self):
        """Test text extraction from different document elements."""
        parser = DOCXParser()
        content = create_complex_docx()
        
        result = parser.parse(content)
        
        assert "Test Document" in result["content"]
        assert "Regular paragraph content" in result["content"]
        assert "Section 1" in result["content"]
        assert "Cell 0,0" in result["content"]
        assert "Cell 2,2" in result["content"] 